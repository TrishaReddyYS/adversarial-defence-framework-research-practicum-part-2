"""Write evaluation results to Excel and Word documents.

Records each benchmark's execution results in a per-dataset Excel workbook and Word report.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

_PACKAGE_ROOT = Path(__file__).resolve().parents[3]
_RESULTS_DIR = _PACKAGE_ROOT / "results"
# Fixed, neutral document timestamp so reports never embed the actual run time.
_NEUTRAL_TIMESTAMP = datetime(2024, 1, 1, 0, 0, 0)


def _ensure_results_dir(out_dir: Path | None) -> Path:
    target = out_dir or _RESULTS_DIR
    target.mkdir(parents=True, exist_ok=True)
    return target


def write_excel(
    rows: Sequence[dict[str, Any]],
    name: str = "evaluation",
    summary: dict[str, Any] | None = None,
    out_dir: Path | None = None,
) -> Path:
    """Write rows (and an optional summary sheet) to the dataset's .xlsx workbook."""
    from openpyxl import Workbook

    target = _ensure_results_dir(out_dir)
    path = target / f"{name}.xlsx"

    wb = Workbook()
    # Keep the workbook metadata neutral so no run timestamp is recorded.
    wb.properties.creator = "adf"
    wb.properties.lastModifiedBy = "adf"
    wb.properties.created = _NEUTRAL_TIMESTAMP
    wb.properties.modified = _NEUTRAL_TIMESTAMP
    ws = wb.active
    ws.title = "results"
    if rows:
        headers = list(rows[0].keys())
        ws.append(headers)
        for row in rows:
            ws.append([row.get(h, "") for h in headers])

    if summary:
        ss = wb.create_sheet("summary")
        ss.append(["metric", "value"])
        for key, value in summary.items():
            ss.append([key, value])

    wb.save(path)
    return path


def write_word(
    rows: Sequence[dict[str, Any]],
    name: str = "evaluation",
    title: str = "Evaluation Results",
    summary: dict[str, Any] | None = None,
    out_dir: Path | None = None,
) -> Path:
    """Write a summary + results table to the dataset's .docx report."""
    from docx import Document

    target = _ensure_results_dir(out_dir)
    path = target / f"{name}.docx"

    doc = Document()
    # Keep the document metadata neutral so no run timestamp is recorded.
    doc.core_properties.author = "adf"
    doc.core_properties.last_modified_by = "adf"
    doc.core_properties.created = _NEUTRAL_TIMESTAMP
    doc.core_properties.modified = _NEUTRAL_TIMESTAMP
    doc.add_heading(title, level=1)

    if summary:
        doc.add_heading("Summary", level=2)
        for key, value in summary.items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    if rows:
        doc.add_heading("Per-item results", level=2)
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for row in rows:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row.get(h, ""))

    doc.save(path)
    return path


def write_all(
    rows: Iterable[dict[str, Any]],
    name: str = "evaluation",
    title: str = "Evaluation Results",
    summary: dict[str, Any] | None = None,
    out_dir: Path | None = None,
) -> tuple[Path, Path]:
    """Write both Excel and Word reports. Returns (xlsx_path, docx_path)."""
    rows = list(rows)
    xlsx = write_excel(rows, name=name, summary=summary, out_dir=out_dir)
    docx = write_word(rows, name=name, title=title, summary=summary, out_dir=out_dir)
    return xlsx, docx
